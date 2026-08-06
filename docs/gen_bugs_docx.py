from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

styles = doc.styles["Normal"]
styles.font.name = "Calibri"
styles.font.size = Pt(11)

title = doc.add_heading("Painel RH — Bugs Encontrados e Corrigidos", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Relatório da revisão de código do projeto painel-rh-intra")
run.italic = True

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    "Este documento lista os 14 bugs identificados em três rodadas de revisão "
    "do código do painel de RH e o que cada correção passou a garantir. "
    "As descrições estão em linguagem funcional — o detalhamento técnico e os "
    "trechos de código correspondentes ficam no documento "
).font.size = Pt(11)
intro.add_run("mudancas-aplicadas.md").italic = True
intro.add_run(".")

doc.add_paragraph()

def bug(numero, titulo, severidade, o_que_acontecia, o_que_foi_corrigido):
    h = doc.add_heading(f"Bug {numero} — {titulo}", level=1)
    p = doc.add_paragraph()
    r = p.add_run("Severidade: ")
    r.bold = True
    p.add_run(severidade)

    p = doc.add_paragraph()
    r = p.add_run("O que acontecia: ")
    r.bold = True
    p.add_run(o_que_acontecia)

    p = doc.add_paragraph()
    r = p.add_run("O que foi corrigido: ")
    r.bold = True
    p.add_run(o_que_foi_corrigido)

    doc.add_paragraph()


bug(
    1,
    "Cálculo de férias ignorava abonos anteriores",
    "Alto",
    "Ao pedir novas férias, o sistema somava apenas os dias de gozo já usufruídos no período aquisitivo — os dias vendidos como abono pecuniário ficavam de fora dessa conta. Um colaborador que já tivesse esgotado seus 30 dias legais usando 20 de gozo mais 10 de abono conseguia abrir uma nova solicitação no mesmo período aquisitivo sem ser barrado pelo sistema. Outro módulo do projeto (o cálculo de prazos de vencimento) já fazia a conta certa; os dois estavam divergentes.",
    "O motor de fatos agora soma gozo e abono de todas as solicitações aprovadas do período, alinhando o resultado com o módulo de prazos. Novas solicitações que estouram o limite legal de 30 dias são bloqueadas com a devida citação do artigo 130 da CLT.",
)

bug(
    2,
    "Aviso segmentado por localidade poderia vazar para toda a região metropolitana",
    "Médio",
    "Ao criar um comunicado direcionado para uma localidade, se o formulário chegasse ao servidor com o valor da localidade em branco, o sistema silenciosamente selecionava todos os colaboradores da Região Metropolitana de Curitiba. Outros tipos de audiência (setor, papel, pessoa) já tinham a proteção correta para essa situação — só a filtragem por localidade estava desprotegida.",
    "O tratamento de audiência por localidade passou a exigir explicitamente um valor. Sem valor, a lista de destinatários vem vazia e nenhum aviso é enviado, mesmo comportamento das demais audiências.",
)

bug(
    3,
    "Envio de notificações em lote era lento e podia estourar timeout",
    "Médio (performance)",
    "Cada notificação individual consultava o banco três vezes só para descobrir quais canais estavam ligados. E o envio para vários destinatários acontecia em série — uma pessoa por vez, esperando WhatsApp e Discord de cada envio antes de começar o próximo. Em cenários realistas (dezenas de solicitações e alguns administradores), o cron de lembretes podia estourar o tempo limite do servidor e ficar preso no meio.",
    "As consultas de configuração foram consolidadas em uma única leitura por tipo de notificação, e o envio em lote passou a acontecer em paralelo, com falhas isoladas por destinatário — se o WhatsApp de alguém der erro, os demais continuam.",
)

bug(
    4,
    "Máscara do RG expunha demais em números curtos",
    "Médio (privacidade)",
    "A regra prometia mostrar somente os três últimos caracteres do RG. Para números curtos, essa regra não estava sendo cumprida: um RG de quatro dígitos era exibido como três asteriscos seguidos dos três últimos caracteres, revelando três dos quatro dígitos.",
    "A máscara agora esconde por completo qualquer RG com até seis caracteres. Só a partir do sétimo caractere o final é exibido, respeitando a promessa da regra.",
)

bug(
    5,
    "Referências entre tabelas do banco estavam declaradas apenas no código",
    "Médio (integridade de dados)",
    "Seis colunas do banco (o gestor de cada colaborador, quem pagou uma solicitação, quem registrou o recibo, quem cancelou, quem aprovou pelo RH, quem aprovou pelo gestor) apontavam para a tabela de usuários apenas por convenção — o banco de dados não tinha nenhuma restrição real dizendo que aqueles identificadores precisavam existir. Outras colunas equivalentes no mesmo esquema já tinham essa proteção; essas seis não.",
    "As referências passaram a existir também no banco, com regra de que, se o usuário referenciado for removido, o campo fica em branco em vez de apontar para nada. Precisa gerar e aplicar uma migração antes de subir para produção — isso está no checklist do documento de mudanças técnicas.",
)

bug(
    6,
    "Antecipação do décimo terceiro podia ser pedida mais de uma vez no mesmo ano",
    "Médio (regra legal)",
    "A Lei 4.749/65 permite apenas uma antecipação da primeira parcela do décimo terceiro por ano-calendário. O sistema aceitava a marcação em cada nova solicitação, sem verificar se o colaborador já havia feito outra antecipação naquele mesmo ano.",
    "Ao pedir nova antecipação, o sistema agora verifica se já existe outra do mesmo ano-calendário. Se existir, um alerta é gerado explicando o artigo 4º da lei, para que o RH e o gestor confiram antes de aprovar. Foi implementado como alerta e não como bloqueio duro porque casos legítimos (correção após cancelamento) podem exigir uma segunda marcação.",
)

bug(
    7,
    "Cálculo de \"hoje\" saía errado após 21h",
    "Baixo",
    "Duas telas — o controle operacional de férias e o calendário anual — calculavam a data de hoje usando o horário universal em vez do horário local. Como o servidor roda em fuso zero, entre 21h de Brasília e a meia-noite o sistema considerava \"hoje\" o dia seguinte. Isso derrubava contadores como \"quantos dias faltam para o pagamento\" e a marcação do dia atual no calendário, em uma janela de três horas por dia.",
    "Foi criado um utilitário compartilhado que devolve a data de hoje sempre no fuso de São Paulo. Ambas as telas passaram a usá-lo.",
)

bug(
    8,
    "Telefones do DDD 55 eram rejeitados pelo WhatsApp",
    "Baixo",
    "A regra de normalização do telefone assumia que qualquer número começando com 55 já vinha com código do país. Isso ignora que 55 também é o DDD da região de Santa Maria, no Rio Grande do Sul. Um colaborador de Santa Maria com número local acabava sendo classificado como inválido e o WhatsApp não era enviado.",
    "A regra passou a considerar o tamanho total do número, e não apenas o começo, para decidir se já vem com código do país. Números de 10 ou 11 dígitos recebem o código do país; números de 12 ou 13 já têm código; qualquer outro tamanho é rejeitado. DDD 55 agora funciona.",
)

bug(
    9,
    "Ponte de acesso ao banco podia quebrar em métodos internos do Drizzle",
    "Baixo",
    "A camada preguiçosa que constrói a conexão com o banco na primeira consulta usava um padrão que pode falhar se a biblioteca de banco de dados usar campos privados de classe internamente. Em produção não estava quebrando, mas o padrão é conhecido por ser frágil e propenso a essa categoria de erro em atualizações da biblioteca.",
    "O acesso passou a garantir que os métodos do banco enxergam o objeto real da conexão em vez da ponte, e métodos são vinculados ao objeto correto. O comportamento visível continua o mesmo e a inicialização preguiçosa foi preservada.",
)

bug(
    10,
    "Contagem de notificações enviadas perdeu significado após paralelização",
    "Médio (regressão introduzida no Bug 3)",
    "Quando as notificações passaram a ser enviadas em paralelo (correção do Bug 3), o campo de resumo \"notificações enviadas\" passou a contar cada tentativa que terminou sem lançar exceção. Como a função de envio nunca lança — ela captura erros internamente e devolve um resultado marcando o canal como falho — a contagem virou matematicamente idêntica ao total de tentativas, perdendo o significado de \"quantas realmente saíram\".",
    "A contagem agora inspeciona o resultado de cada envio e considera \"notificado\" quem teve a notificação gravada dentro do sistema ou pelo menos um canal externo entregue. O número volta a refletir entregas reais.",
)

bug(
    11,
    "Segunda execução da migração era abortada pela trava de segurança",
    "Médio (bloqueia deploy)",
    "O script de migração tem uma trava que aborta se detectar tabelas que não pertencem a este projeto — proteção contra rodar migrações no banco errado. A lista de tabelas conhecidas estava desatualizada: faltava a tabela de configurações de notificações. Consequência: a primeira migração criava essa tabela normalmente, mas qualquer execução seguinte abortava reclamando que a tabela \"não é deste projeto\", travando o deploy.",
    "A lista foi atualizada para incluir a tabela ausente. A trava continua ativa e protegendo contra rodar no banco errado — só agora reconhece corretamente as tabelas do próprio projeto.",
)

bug(
    12,
    "Motor de decisão de férias também usava horário universal",
    "Médio",
    "Mesma classe do Bug 7, em outro arquivo. O motor que decide se uma solicitação de férias pode ser aprovada ou não usava horário universal para saber \"que dia é hoje\". Isso significa que, entre 21h e meia-noite de Brasília, o sistema considerava que a solicitação começava no passado (marcando erroneamente como fora do prazo) e subtraía um dia inteiro do cálculo de \"quantos dias de antecedência a pessoa avisou\".",
    "O motor passou a usar o mesmo utilitário de \"hoje no fuso de São Paulo\" já usado nas telas de controle e no calendário. Bloqueios e alertas noturnos deixaram de ser distorcidos por três horas de defasagem.",
)

bug(
    13,
    "Relatórios em CSV eram vulneráveis a injeção de fórmula",
    "Médio",
    "Os dois relatórios exportados em CSV (cadastro de colaboradores e controle de férias) tratavam adequadamente aspas, mas não protegiam contra campos começando com sinais de igual, mais, menos, arroba ou tabulação. Programas como Excel e LibreOffice tratam esses campos como fórmulas, mesmo quando envolvidos por aspas. Um colaborador cadastrado com um nome começando com sinal de igual poderia executar cálculos — ou, no Windows, até comandos do sistema operacional — quando o arquivo fosse aberto pela contabilidade ou pela folha de pagamento.",
    "Foi criado um utilitário central de escape para CSV que, além de tratar aspas, adiciona um apóstrofo no começo de campos suspeitos, neutralizando a interpretação como fórmula. Os dois relatórios passaram a usar esse utilitário.",
)

bug(
    14,
    "Autocompletar de CEP podia sobrescrever endereço com dados errados",
    "Baixo",
    "No formulário de cadastro de colaborador, ao digitar o CEP, o sistema consulta o serviço externo para preencher rua, cidade e estado. A implementação disparava uma consulta a cada CEP válido digitado, sem cancelar consultas anteriores. Se o usuário trocasse o CEP antes da primeira resposta chegar, uma resposta lenta antiga poderia chegar depois e sobrescrever os campos com dados do CEP anterior.",
    "O formulário passou a numerar cada consulta em ordem, descartando qualquer resposta que não seja a mais recente. Somente a resposta do CEP atualmente digitado atualiza os campos de endereço.",
)

doc.add_paragraph()
doc.add_heading("Resumo", level=1)
tabela = doc.add_table(rows=1, cols=3)
tabela.style = "Light Grid Accent 1"
h = tabela.rows[0].cells
h[0].text = "Nº"
h[1].text = "Bug"
h[2].text = "Severidade"

linhas = [
    ("1", "Cálculo de férias ignorava abonos anteriores", "Alto"),
    ("2", "Aviso por localidade sem valor vazava para toda a RMC", "Médio"),
    ("3", "Envio de notificações em lote era lento", "Médio"),
    ("4", "Máscara do RG expunha demais em números curtos", "Médio"),
    ("5", "Referências entre tabelas só no código, não no banco", "Médio"),
    ("6", "Antecipação do 13º sem checagem de unicidade anual", "Médio"),
    ("7", "Cálculo de hoje em UTC nas telas de controle e calendário", "Baixo"),
    ("8", "Telefones do DDD 55 rejeitados", "Baixo"),
    ("9", "Ponte de acesso ao banco frágil a campos privados", "Baixo"),
    ("10", "Contagem de notificações enviadas perdeu significado", "Médio"),
    ("11", "Migração abortada na segunda execução", "Médio"),
    ("12", "Motor de decisão de férias também em UTC", "Médio"),
    ("13", "CSV vulnerável a injeção de fórmula", "Médio"),
    ("14", "Autocompletar de CEP com sobrescrita por resposta antiga", "Baixo"),
]
for l in linhas:
    row = tabela.add_row().cells
    row[0].text = l[0]
    row[1].text = l[1]
    row[2].text = l[2]

doc.add_paragraph()
final = doc.add_paragraph()
final.add_run(
    "Total: 14 bugs identificados e corrigidos — 1 de severidade alta, "
    "9 de severidade média, 4 de severidade baixa. Nenhum de severidade crítica. "
    "A validação automática do projeto (verificação de tipos e análise estática) "
    "continua limpa após todas as correções."
).font.size = Pt(11)

doc.save("D:/painel-rh-intra-main/docs/bugs-corrigidos.docx")
print("OK")
